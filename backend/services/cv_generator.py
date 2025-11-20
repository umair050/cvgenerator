from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import requests
import logging
from typing import List, Dict, Optional
from io import BytesIO

logger = logging.getLogger(__name__)


class CVGenerator:
    """Generate DOCX files from formatted CV text"""
    
    def __init__(self):
        self.heading_styles = {
            'h1': {'size': 16, 'bold': True, 'space_after': 6},
            'h2': {'size': 14, 'bold': True, 'space_after': 6},
            'h3': {'size': 12, 'bold': True, 'space_after': 4},
        }
        # Datamatics logo - try local file first, then URL
        self.logo_path = os.path.join(os.path.dirname(__file__), "..", "assets", "logo.png")
        self.logo_url = "https://datamaticstechnologies.com/wp-content/uploads/2025/09/Default-Logo-Final-for-Black-Background-2-scaled-150x45.png"
    
    async def generate_docx(self, cv_text: str, output_path: str, format_type: str = "modern"):
        """
        Generate a DOCX file from formatted CV text
        """
        if format_type == "datamatics":
            await self._generate_datamatics_format(cv_text, output_path)
        elif format_type == "two-column" or format_type == "modern":
            await self._generate_two_column_format(cv_text, output_path)
        else:
            await self._generate_standard_format(cv_text, output_path)
    
    async def _generate_datamatics_format(self, cv_text: str, output_path: str):
        """
        Generate CV in Datamatics two-column format
        """
        doc = Document()
        
        # Set page margins with more spacing on left and right
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)  # Increased to accommodate header
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.90)  # Increased left margin
            section.right_margin = Inches(0.90)  # Increased right margin
        
        # Set default font - paragraph size is 9pt
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(9)  # Paragraph font size set to 9pt
        
        # Parse structured content
        content = self._parse_datamatics_content(cv_text)
        
        # Add header to the document (Word header section)
        self._add_document_header(doc, content.get('header', ''))
        
        # Create two-column layout using table
        table = doc.add_table(rows=1, cols=2)
        # Remove table borders for cleaner look
        table.style = None
        
        # Set column widths with spacing between columns
        for row in table.rows:
            row.cells[0].width = Inches(3)  # Left column width
            row.cells[1].width = Inches(3)  # Right column width with gap between
            # Remove borders from cells and add spacing between columns
            for idx, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
                
                # Add margin/padding to create spacing between columns
                tcMar = OxmlElement('w:tcMar')
                # Add right margin to left cell for spacing
                if idx == 0:
                    right_margin = OxmlElement('w:right')
                    right_margin.set(qn('w:w'), '360')  # 360 twips = ~0.25 inches spacing
                    right_margin.set(qn('w:type'), 'dxa')
                    tcMar.append(right_margin)
                # Add left margin to right cell for spacing
                elif idx == 1:
                    left_margin = OxmlElement('w:left')
                    left_margin.set(qn('w:w'), '360')  # 360 twips = ~0.25 inches spacing
                    left_margin.set(qn('w:type'), 'dxa')
                    tcMar.append(left_margin)
                tcPr.append(tcMar)
        
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]
        
        # LEFT COLUMN
        left_para = left_cell.paragraphs[0]
        left_para.clear()
        
        # Technical Skills - with categories
        if content.get('technical_skills'):
            # First section - header formatting with 10pt font
            tech_header = left_cell.add_paragraph()
            tech_run = tech_header.add_run("Technical Skills")
            tech_run.font.bold = True
            tech_run.font.size = Pt(10)  # Header font size set to 10pt
            tech_run.font.name = 'Calibri'
            tech_header.paragraph_format.space_before = Pt(0)
            tech_header.paragraph_format.space_after = Pt(3)  # Tighter spacing after header
            
            # Process skills with categories - format: "Category: skill1, skill2, skill3"
            skills_text = content['technical_skills']
            if isinstance(skills_text, str):
                # Split by newlines to get categories
                skill_lines = [s.strip() for s in skills_text.split('\n') if s.strip()]
            else:
                skill_lines = skills_text if isinstance(skills_text, list) else []
            
            for line in skill_lines:
                # Check if line contains category format (has colon)
                if ':' in line:
                    # Category format: "Category Name: skill1, skill2, skill3"
                    parts = line.split(':', 1)
                    category = parts[0].strip()
                    skills = parts[1].strip() if len(parts) > 1 else ""
                    
                    # Clean up skills text - remove any markdown bold formatting
                    skills = re.sub(r'\*\*([^*]+)\*\*', r'\1', skills)
                    skills = re.sub(r'\*([^*]+)\*', r'\1', skills)
                    
                    # Add category with skills
                    if category and skills:
                        self._add_category_skill(left_cell, category, skills)
                else:
                    # Fallback: treat as regular bullet point if no category
                    # Clean up any markdown formatting
                    line_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
                    line_clean = re.sub(r'\*([^*]+)\*', r'\1', line_clean)
                    self._add_bullet_point(left_cell, line_clean)
            
            self._add_spacing(left_cell)
        
        # Functional Skills
        if content.get('functional_skills'):
            self._add_section_header(left_cell, "Functional Skills")
            skills_text = content['functional_skills']
            if isinstance(skills_text, str):
                skills_list = [s.strip() for s in skills_text.replace(',', '\n').split('\n') if s.strip()]
            else:
                skills_list = skills_text if isinstance(skills_text, list) else []
            for skill in skills_list:
                self._add_bullet_point(left_cell, skill)  # Already justified in _add_bullet_point
            self._add_spacing(left_cell)
        
        # Industry Experience
        if content.get('industry_experience'):
            self._add_section_header(left_cell, "Industry Experience")
            industry_text = content['industry_experience']
            if isinstance(industry_text, str):
                industry_list = [i.strip() for i in industry_text.replace(',', '\n').split('\n') if i.strip()]
            else:
                industry_list = industry_text if isinstance(industry_text, list) else []
            for industry in industry_list:
                self._add_bullet_point(left_cell, industry)  # Already justified in _add_bullet_point
            self._add_spacing(left_cell)
        
        # RIGHT COLUMN
        right_para = right_cell.paragraphs[0]
        right_para.clear()
        
        # Summary
        if content.get('summary'):
            # Summary header - first section in right column with 10pt font
            summary_header = right_cell.add_paragraph()
            summary_run = summary_header.add_run("Summary")
            summary_run.font.bold = True
            summary_run.font.size = Pt(10)  # Header font size set to 10pt
            summary_run.font.name = 'Calibri'
            summary_header.paragraph_format.space_before = Pt(0)
            summary_header.paragraph_format.space_after = Pt(3)
            # Summary paragraphs - regular text, not bullets - left-aligned
            summary_text = content['summary']
            if isinstance(summary_text, str):
                # Split by double newlines for paragraphs
                paragraphs = [p.strip() for p in summary_text.split('\n\n') if p.strip()]
                if not paragraphs:
                    # Fallback: split by single newlines
                    paragraphs = [p.strip() for p in summary_text.split('\n') if p.strip()]
            else:
                paragraphs = summary_text if isinstance(summary_text, list) else [str(summary_text)]
            
            for para in paragraphs:
                p = right_cell.add_paragraph(para)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)  # Compact spacing after summary paragraphs
                p.paragraph_format.left_indent = Inches(0.1)  # Minimal left indent
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # Exact spacing - no extra space on wrapped lines
                p.paragraph_format.line_spacing = Pt(11)  # 11pt line spacing for 9pt font (tight but readable)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justified text
                for run in p.runs:
                    run.font.size = Pt(9)  # Paragraph font size set to 9pt
                    run.font.name = 'Calibri'
                    run.font.bold = False
            self._add_spacing(right_cell)
        
        # Education/Qualifications/Certifications - Combined section below Summary
        has_certs = content.get('certifications')
        has_edu = content.get('education')
        
        if has_certs or has_edu:
            self._add_section_header(right_cell, "Education/Qualifications/Certifications")
            
            # First, list all certifications
            if has_certs:
                cert_text = content['certifications']
                if isinstance(cert_text, str):
                    cert_list = [c.strip() for c in cert_text.replace(',', '\n').split('\n') if c.strip()]
                else:
                    cert_list = cert_text if isinstance(cert_text, list) else []
                
                for cert in cert_list:
                    # Clean up certification entry
                    cert_clean = cert.strip()
                    # Remove any markdown formatting
                    cert_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', cert_clean)
                    cert_clean = re.sub(r'^\d+\.\s*', '', cert_clean)  # Remove numbering
                    if cert_clean and len(cert_clean) > 3:
                        self._add_bullet_point(right_cell, cert_clean)
            
            # Then, add education (only most recent)
            if has_edu:
                edu_text = content['education']
                if isinstance(edu_text, str):
                    # Split by newlines or commas, but handle "Degree | Institution | Year" format
                    edu_list = [e.strip() for e in edu_text.replace(',', '\n').split('\n') if e.strip()]
                else:
                    edu_list = edu_text if isinstance(edu_text, list) else []
                
                # Only use the LAST/MOST RECENT education entry
                if edu_list:
                    # Take only the last entry
                    edu = edu_list[-1]
                    # Clean up the education entry
                    edu_clean = edu.strip()
                    # Remove any markdown formatting
                    edu_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', edu_clean)
                    edu_clean = re.sub(r'^\d+\.\s*', '', edu_clean)  # Remove numbering
                    
                    # Format: "Degree | Institution | Year" or "Degree | Institution"
                    # Ensure proper pipe separation if not present
                    if '|' not in edu_clean and len(edu_clean) > 10:
                        # Try to detect pattern: Degree Institution Year or Degree Institution
                        # If it looks like "Bachelor of Science in Computer Science Mohammad Ali Jinnah University 2016"
                        # This is already a complex case, just use as-is
                        pass
                    
                    if edu_clean and len(edu_clean) > 5:
                        self._add_bullet_point_bold(right_cell, edu_clean)  # Use bold for education
            
            self._add_spacing(right_cell)
        
        # Work Experience - Add as new row in table with merged cell spanning both columns
        if content.get('projects'):
            # Add a new row to the table for Work Experience
            work_exp_row = table.add_row()
            # Merge the cells to span both columns
            work_exp_cell = work_exp_row.cells[0]
            work_exp_row.cells[1].merge(work_exp_cell)
            
            # Remove borders from merged cell
            tc = work_exp_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)
            
            # Add spacing before work experience section
            spacing_before = work_exp_cell.add_paragraph()
            spacing_before.paragraph_format.space_after = Pt(8)  # More space before section
            
            work_exp_heading = work_exp_cell.add_paragraph()
            work_exp_run = work_exp_heading.add_run("Work Experience")
            work_exp_run.font.size = Pt(10)  # Header font size set to 10pt
            work_exp_run.font.bold = True
            work_exp_run.font.name = 'Calibri'
            # Set red color for "Work Experience" heading
            try:
                work_exp_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
            except:
                # If RGBColor doesn't work, use alternative method
                pass
            work_exp_heading.paragraph_format.space_before = Pt(0)
            work_exp_heading.paragraph_format.space_after = Pt(8)  # More spacing after heading for better separation
            
            for project in content['projects']:
                # Skip empty projects
                if not project.get('title') and not project.get('responsibilities'):
                    continue
                    
                # Project header: Match exact format from samples
                # Format options:
                # 1. "JOB TITLE – COMPANY | DATES | LOCATION" (en-dash format)
                # 2. "JOB TITLE, COMPANY, LOCATION | DATES" (comma-separated format)
                header_p = work_exp_cell.add_paragraph()
                title_text = project.get('title', '') or 'Position / Company'
                location = project.get('location', '').strip() if project.get('location') else ''
                
                # Build complete header line - all uppercase, all bold
                header_line_parts = []
                
                # Check if title contains " / " (forward slash) - use en-dash format
                if ' / ' in title_text:
                    parts = title_text.split(' / ', 1)
                    job_title = parts[0].strip().upper()
                    company = parts[1].strip().upper() if len(parts) > 1 else ''
                    
                    # Format: "JOB TITLE – COMPANY"
                    header_line_parts.append(job_title)
                    if company:
                        header_line_parts.append(' – ')
                        header_line_parts.append(company)
                # Check if title contains comma - use comma-separated format
                elif ',' in title_text and not location:
                    # Format: "JOB TITLE, COMPANY, LOCATION | DATES"
                    # Split by comma to get title, company, and possibly location
                    parts = [p.strip() for p in title_text.split(',')]
                    if len(parts) >= 1:
                        header_line_parts.append(parts[0].upper())
                    if len(parts) >= 2:
                        header_line_parts.append(', ')
                        header_line_parts.append(parts[1].upper())
                    if len(parts) >= 3:
                        # Third part might be location
                        header_line_parts.append(', ')
                        header_line_parts.append(parts[2].upper())
                else:
                    # Simple format - just title
                    header_line_parts.append(title_text.upper())
                
                # Add dates (if present) - format: " | DATES"
                if project.get('dates'):
                    dates_text = project['dates']
                    # Remove "Dates: " prefix if present
                    dates_text = re.sub(r'^Dates:\s*', '', dates_text, flags=re.IGNORECASE)
                    # Normalize date format - handle both "MM/YYYY - MM/YYYY" and "YYYY-PRESENT"
                    # Keep hyphens as-is for dates (don't convert to en-dash for date ranges)
                    dates_text = dates_text.strip()
                    # Convert to uppercase
                    dates_text = dates_text.upper()
                    header_line_parts.append(' | ')
                    header_line_parts.append(dates_text)
                    
                    # Add location (if present and not already in title) - format: " | LOCATION"
                    if location and location not in title_text.upper():
                        location_text = location.upper()
                        header_line_parts.append(' | ')
                        header_line_parts.append(location_text)
                
                # Add all parts as bold with 10pt header font
                for part in header_line_parts:
                    run = header_p.add_run(part)
                    run.font.bold = True
                    run.font.size = Pt(10)  # Header font size set to 10pt
                    run.font.name = 'Calibri'
                
                header_p.paragraph_format.space_before = Pt(0)
                header_p.paragraph_format.space_after = Pt(6)  # Spacing after job title before first category (matches image)
                header_p.paragraph_format.left_indent = Inches(0)  # Job title at left edge
                
                # Responsibilities - with optional category headers
                if project.get('responsibilities'):
                    for resp in project['responsibilities']:
                        # Clean up responsibility text
                        resp_clean = resp.strip()
                        
                        # Check if this is a category header (ends with ":" and doesn't start with bullet)
                        is_category_header = (
                            resp_clean.endswith(':') and 
                            not resp_clean.startswith('-') and 
                            not resp_clean.startswith('•') and 
                            not resp_clean.startswith('▪') and
                            len(resp_clean) < 100 and  # Headers are typically shorter
                            resp_clean.count(':') == 1 and  # Only one colon at the end
                            not resp_clean.startswith('Technologies')  # Not the technologies line
                        )
                        
                        if is_category_header:
                            # Render as category header (bold, no bullet) - matching image format
                            category_p = work_exp_cell.add_paragraph()
                            category_run = category_p.add_run(resp_clean)
                            category_run.font.bold = True
                            category_run.font.size = Pt(9)
                            category_run.font.name = 'Calibri'
                            category_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                            category_p.paragraph_format.space_before = Pt(8)  # Space before category (matches image)
                            category_p.paragraph_format.space_after = Pt(3)  # Space after category before bullets
                            category_p.paragraph_format.left_indent = Inches(0)  # No indent - flush left (matches image)
                            category_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            category_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # Exact spacing
                            category_p.paragraph_format.line_spacing = Pt(11)  # Tight spacing
                        else:
                            # Regular responsibility bullet point
                            # Remove any leading/trailing punctuation issues
                            resp_clean = re.sub(r'^[-•*▪]\s*', '', resp_clean).strip()
                            # Remove any markdown bold formatting
                            resp_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', resp_clean)
                            resp_clean = re.sub(r'\*([^*]+)\*', r'\1', resp_clean)
                            # Remove any trailing periods that might interfere
                            resp_clean = resp_clean.rstrip('.')
                            if resp_clean and len(resp_clean) > 3:  # Ensure meaningful content
                                # Use bullet point function which ensures normal text (not bold)
                                self._add_bullet_point(work_exp_cell, resp_clean)
                
                # Technologies (if present) - justified format
                if project.get('technologies'):
                    tech_p = work_exp_cell.add_paragraph()
                    tech_run = tech_p.add_run('Technologies: ')
                    tech_run.font.italic = False  # Not italic - plain text
                    tech_run.font.bold = False
                    tech_run.font.size = Pt(9)  # Paragraph font size set to 9pt
                    tech_run.font.name = 'Calibri'
                    tech_text = tech_p.add_run(', '.join(project['technologies']))
                    tech_text.font.size = Pt(9)  # Paragraph font size set to 9pt
                    tech_text.font.name = 'Calibri'
                    tech_text.font.bold = False
                    tech_p.paragraph_format.space_before = Pt(3)  # Small space before technologies
                    tech_p.paragraph_format.space_after = Pt(3)  # Spacing after technologies
                    tech_p.paragraph_format.left_indent = Inches(0)  # No indent - flush left
                    tech_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # Exact spacing
                    tech_p.paragraph_format.line_spacing = Pt(11)  # Tight spacing
                    tech_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justified text
                
                # Add spacing between projects (if not last one and project has content)
                if project != content['projects'][-1] and (project.get('responsibilities') or project.get('technologies')):
                    spacing_p = work_exp_cell.add_paragraph()
                    spacing_p.paragraph_format.space_after = Pt(10)  # More spacing between different job entries (matches image)
        
        doc.save(output_path)
    
    def _add_section_header(self, cell, text: str):
        """Add a section header to a cell with 10pt font and left alignment for headers"""
        p = cell.add_paragraph()
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(10)  # Header font size set to 10pt
        run.font.name = 'Calibri'
        # Match PDF spacing exactly - consistent spacing between sections
        p.paragraph_format.space_before = Pt(6)  # Slightly less space before
        p.paragraph_format.space_after = Pt(3)  # Tighter spacing after header
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.first_line_indent = Inches(0)
        # Headers are left-aligned (not justified) for clean section headers
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_category_skill(self, cell, category: str, skills: str):
        """Add a category with comma-separated skills using Word's native bullet style"""
        p = cell.add_paragraph()
        
        # Apply Word's List Bullet style
        p.style = 'List Bullet'
        
        # Clean category name - remove any markdown bold formatting
        category_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', category)
        category_clean = re.sub(r'\*([^*]+)\*', r'\1', category_clean)
        
        # Add category name in bold
        category_run = p.add_run(f"{category_clean}: ")
        category_run.font.bold = True
        category_run.font.size = Pt(9)  # Paragraph font size set to 9pt
        category_run.font.name = 'Calibri'
        
        # Add skills (comma-separated) in regular font (NOT bold)
        skills_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', skills)
        skills_clean = re.sub(r'\*([^*]+)\*', r'\1', skills_clean)
        skills_run = p.add_run(skills_clean)
        skills_run.font.bold = False
        skills_run.font.size = Pt(9)  # Paragraph font size set to 9pt
        skills_run.font.name = 'Calibri'
        
        # Format for compact layout
        p.paragraph_format.left_indent = Inches(0.2)  # Minimal left indent (compact format)
        p.paragraph_format.space_after = Pt(0)  # No space after for compact look
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # Exact spacing
        p.paragraph_format.line_spacing = Pt(11)  # Tight spacing
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left-aligned
    
    def _add_bullet_point(self, cell, text: str):
        """Add a bullet point using Word's native bullet list style"""
        p = cell.add_paragraph()
        
        # Apply Word's List Bullet style to make bullets recognized by Word
        p.style = 'List Bullet'
        
        # Clean text thoroughly - remove any markdown bold formatting
        text_clean = text
        # Remove markdown bold (**text** or *text*)
        text_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', text_clean)
        text_clean = re.sub(r'\*([^*]+)\*', r'\1', text_clean)
        # Remove any HTML bold tags if present
        text_clean = re.sub(r'<b>|</b>|<strong>|</strong>', '', text_clean, flags=re.IGNORECASE)
        # Remove any leading/trailing whitespace
        text_clean = text_clean.strip()
        
        # Add text with proper formatting
        text_run = p.add_run(text_clean)
        text_run.font.size = Pt(9)  # Paragraph font size set to 9pt
        text_run.font.name = 'Calibri'
        text_run.font.bold = False  # Not bold
        text_run.font.italic = False
        
        # Format the paragraph for compact, justified layout (work experience)
        p.paragraph_format.left_indent = Inches(0.25)  # Bullet indent (matches image)
        p.paragraph_format.first_line_indent = Inches(-0.15)  # Hanging indent for bullet (matches image)
        p.paragraph_format.space_after = Pt(3)  # Small space after for readability (matches image)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # Exact spacing - no extra space on wrapped lines
        p.paragraph_format.line_spacing = Pt(11)  # 11pt line spacing for 9pt font (tight but readable)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justified text
    
    def _add_bullet_point_bold(self, cell, text: str):
        """Add a bullet point using Word's native bullet list style with BOLD text"""
        p = cell.add_paragraph()
        
        # Apply Word's List Bullet style to make bullets recognized by Word
        p.style = 'List Bullet'
        
        # Clean text thoroughly - remove any markdown formatting
        text_clean = text
        text_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', text_clean)
        text_clean = re.sub(r'\*([^*]+)\*', r'\1', text_clean)
        text_clean = re.sub(r'<b>|</b>|<strong>|</strong>', '', text_clean, flags=re.IGNORECASE)
        text_clean = text_clean.strip()
        
        # Add text with BOLD formatting
        text_run = p.add_run(text_clean)
        text_run.font.size = Pt(9)  # Paragraph font size set to 9pt
        text_run.font.name = 'Calibri'
        text_run.font.bold = True  # Make education text BOLD
        text_run.font.italic = False
        
        # Format the paragraph for compact, justified layout
        p.paragraph_format.left_indent = Inches(0.3)  # Minimal left indent
        p.paragraph_format.space_after = Pt(2)  # Small space after for readability
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # Exact spacing - no extra space on wrapped lines
        p.paragraph_format.line_spacing = Pt(11)  # 11pt line spacing for 9pt font (tight but readable)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justified text
    
    def _add_bullet_point_doc(self, doc, text: str):
        """Add a bullet point to document using Word's native bullet style"""
        p = doc.add_paragraph()
        
        # Apply Word's List Bullet style
        p.style = 'List Bullet'
        
        # Add the text with 9pt font
        text_run = p.add_run(text)
        text_run.font.size = Pt(9)  # Paragraph font size set to 9pt
        text_run.font.name = 'Calibri'
        text_run.font.bold = False
        
        # Format for compact layout
        p.paragraph_format.left_indent = Inches(0.2)  # Compact indent
        p.paragraph_format.space_after = Pt(0)  # No space after for compact look
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # Exact spacing
        p.paragraph_format.line_spacing = Pt(11)  # Tight spacing
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_spacing(self, cell):
        """Add spacing paragraph"""
        cell.add_paragraph()
    
    def _add_document_header(self, doc, header_text: str):
        """
        Add header section to the Word document (first page only)
        Header contains: Name | Position on left, Logo on right
        """
        # Access the header of the first section (default section)
        section = doc.sections[0]
        
        # Enable different first page header
        section.different_first_page_header_footer = True
        
        # Get the first page header
        header = section.first_page_header
        
        # Create a table in the header with 2 columns
        header_table = header.add_table(rows=1, cols=2, width=Inches(7))
        header_table.style = None
        
        # Set column widths
        header_table.columns[0].width = Inches(4.5)  # Left: Name/Position
        header_table.columns[1].width = Inches(2.5)  # Right: Logo
        
        # Remove borders from header table
        for row in header_table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
        
        # Left cell: Name and Position
        left_cell = header_table.rows[0].cells[0]
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        if header_text:
            # Parse "Name | Position" format
            if '|' in header_text:
                parts = [p.strip() for p in header_text.split('|')]
                name = parts[0] if parts else ""
                position = parts[1] if len(parts) > 1 else ""
            else:
                name = header_text
                position = ""
            
            # Create a single line with "Name | Position"
            if name and position:
                header_p = left_cell.add_paragraph()
                header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Name (14pt, bold, black)
                name_run = header_p.add_run(name + " | ")
                name_run.font.size = Pt(14)
                name_run.font.bold = True
                name_run.font.name = 'Calibri'
                name_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                
                # Position (14pt, bold, blue)
                pos_run = header_p.add_run(position)
                pos_run.font.size = Pt(14)
                pos_run.font.bold = True
                pos_run.font.name = 'Calibri'
                pos_run.font.color.rgb = RGBColor(91, 155, 213)  # Blue color
                
                header_p.paragraph_format.space_before = Pt(0)
                header_p.paragraph_format.space_after = Pt(0)
            elif name:
                header_p = left_cell.add_paragraph()
                header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                name_run = header_p.add_run(name)
                name_run.font.size = Pt(14)
                name_run.font.bold = True
                name_run.font.name = 'Calibri'
                name_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                header_p.paragraph_format.space_before = Pt(0)
                header_p.paragraph_format.space_after = Pt(0)
        
        # Right cell: Logo
        right_cell = header_table.rows[0].cells[1]
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Add logo
        try:
            logo_image = self._download_logo()
            if logo_image:
                logo_para = right_cell.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = logo_para.add_run()
                # Logo width: match the image size
                run.add_picture(logo_image, width=Inches(1.8))
                logo_para.paragraph_format.space_before = Pt(0)
                logo_para.paragraph_format.space_after = Pt(0)
        except Exception as e:
            logger.warning(f"Could not add logo to header: {str(e)}")
            # Continue without logo if download fails
        
        # Set header margin (increased for larger font)
        section.header_distance = Inches(0.4)
    
    def _download_logo(self) -> Optional[BytesIO]:
        """Get Datamatics logo from local file or download from URL"""
        # Try local file first
        if os.path.exists(self.logo_path):
            try:
                with open(self.logo_path, 'rb') as f:
                    return BytesIO(f.read())
            except Exception as e:
                logger.warning(f"Failed to read local logo: {str(e)}")
        
        # Fallback to downloading from URL
        try:
            logger.info(f"Downloading logo from URL: {self.logo_url}")
            response = requests.get(self.logo_url, timeout=10)
            if response.status_code == 200:
                logger.info("Logo downloaded successfully")
                return BytesIO(response.content)
            else:
                logger.warning(f"Logo download failed with status code: {response.status_code}")
        except Exception as e:
            logger.warning(f"Failed to download logo: {str(e)}")
        
        logger.warning("No logo available - continuing without logo")
        return None
    
    def _parse_datamatics_content(self, text: str) -> Dict:
        """Parse structured content from LLM output"""
        content = {
            'header': '',
            'technical_skills': [],
            'industry_experience': [],
            'functional_skills': [],
            'education': [],
            'summary': [],
            'certifications': [],
            'projects': []
        }
        
        lines = text.split('\n')
        current_section = None
        current_project = None
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            line_lower = line.lower()
            
            # Header detection (Name | Title format)
            if '|' in line and not current_section and not content['header'] and len(line) < 100:
                # Check if it looks like a header (not a bullet point or section)
                if not line.startswith('-') and not line.startswith('•') and not line.startswith('['):
                    content['header'] = line
                    i += 1
                    continue
            
            # Section detection
            if '[LEFT_COLUMN_START]' in line or '[LEFT_COLUMN]' in line:
                i += 1
                continue
            if '[LEFT_COLUMN_END]' in line or '[RIGHT_COLUMN_START]' in line:
                current_section = None
                i += 1
                continue
            if '[RIGHT_COLUMN_END]' in line or '[PROJECTS_EXPERIENCE]' in line:
                current_section = None
                i += 1
                continue
            
            # Detect sections - be more flexible with detection
            if ('technical skills' in line_lower or (line_lower.startswith('technical') and 'skill' in line_lower)) and (':' in line or len(line) < 30):
                current_section = 'technical_skills'
                i += 1
                continue
            elif ('industry experience' in line_lower or (line_lower.startswith('industry') and 'experience' in line_lower)) and (':' in line or len(line) < 30):
                current_section = 'industry_experience'
                i += 1
                continue
            elif ('functional skills' in line_lower or (line_lower.startswith('functional') and 'skill' in line_lower)) and (':' in line or len(line) < 30):
                current_section = 'functional_skills'
                i += 1
                continue
            elif ('education' in line_lower and 'qualification' in line_lower and 'certification' in line_lower) or \
                 ('education/qualification' in line_lower.replace(' ', '')) or \
                 (('education' in line_lower or 'qualification' in line_lower) and 'certification' in line_lower and '/' in line):
                # Combined Education/Qualifications/Certifications section
                current_section = 'education_certifications_combined'
                i += 1
                continue
            elif 'education' in line_lower and ':' in line and 'certification' not in line_lower:
                current_section = 'education'
                i += 1
                continue
            elif 'summary' in line_lower and (':' in line or len(line) < 30):
                current_section = 'summary'
                i += 1
                continue
            elif ('certification' in line_lower or 'trainings' in line_lower) and (':' in line or len(line) < 40):
                current_section = 'certifications'
                i += 1
                continue
            elif 'projects experience' in line_lower or (line_lower.startswith('projects') and 'experience' in line_lower):
                current_section = 'projects'
                i += 1
                continue
            
            # Process content based on current section
            if current_section in ['technical_skills', 'industry_experience', 'functional_skills', 'education', 'certifications', 'education_certifications_combined']:
                # Skip if this line is another section header
                if any(keyword in line_lower for keyword in ['technical skills', 'industry experience', 'functional skills', 'summary', 'projects']):
                    # Check for combined section first (contains both education and certification keywords)
                    if ('education' in line_lower and 'qualification' in line_lower and 'certification' in line_lower):
                        current_section = 'education_certifications_combined'
                    # This might be a new section, reset
                    elif 'technical skills' in line_lower or (line_lower.startswith('technical') and 'skill' in line_lower):
                        current_section = 'technical_skills'
                    elif 'industry experience' in line_lower:
                        current_section = 'industry_experience'
                    elif 'functional skills' in line_lower:
                        current_section = 'functional_skills'
                    elif 'education' in line_lower and 'certification' not in line_lower:
                        current_section = 'education'
                    elif 'certification' in line_lower or 'trainings' in line_lower:
                        current_section = 'certifications'
                    else:
                        current_section = None
                    i += 1
                    continue
                # Remove bullet markers
                clean_line = re.sub(r'^[-•*▪]\s*', '', line).strip()
                
                # For technical_skills, handle category format: "Category: skill1, skill2"
                if current_section == 'technical_skills':
                    # Check if line has category format (contains colon)
                    if ':' in clean_line:
                        # This is a category line, add as-is (preserve category: skills format)
                        if clean_line and len(clean_line) > 3 and clean_line.lower() not in ['technical skills', 'industry experience', 'functional skills', 'education', 'certification', 'certifications', 'summary', 'projects experience']:
                            content[current_section].append(clean_line)
                    else:
                        # Single skill without category, might be from old format
                        # Try to preserve it
                        if clean_line and len(clean_line) > 3 and clean_line.lower() not in ['technical skills', 'industry experience', 'functional skills', 'education', 'certification', 'certifications', 'summary', 'projects experience']:
                            content[current_section].append(clean_line)
                # For combined education/certifications section, separate education from certifications
                elif current_section == 'education_certifications_combined':
                    # Determine if this line is education or certification based on keywords
                    is_education = any(keyword in clean_line.lower() for keyword in [
                        'bachelor', 'master', 'phd', 'doctorate', 'degree', 'diploma', 
                        'b.sc', 'b.tech', 'm.sc', 'm.tech', 'mba', 'bba', 'bca', 'mca',
                        'university', 'college', 'institute', 'b.e', 'be ', 'engineering'
                    ])
                    
                    # Check if this looks like a location/year continuation (e.g., "Hyderabad | 2009")
                    is_location_year = (
                        not is_education and 
                        content.get('education') and  # Only if we already have education
                        ('|' in clean_line or re.search(r'\b(19|20)\d{2}\b', clean_line)) and
                        len(clean_line) < 50 and
                        not any(cert_keyword in clean_line.lower() for cert_keyword in [
                            'aws', 'azure', 'certified', 'certification', 'microsoft', 'oracle',
                            'sap', 'cisco', 'google', 'professional', 'associate', 'expert',
                            'cloud', 'administrator', 'developer', 'architect', 'engineer'
                        ])
                    )
                    
                    if is_education:
                        # This is an education entry
                        edu_clean = clean_line
                        # Remove any trailing periods or extra formatting
                        edu_clean = re.sub(r'\.$', '', edu_clean).strip()
                        if edu_clean and len(edu_clean) > 5:
                            # Only keep the last education entry (replace previous ones)
                            content['education'] = [edu_clean]
                    elif is_location_year:
                        # This is a continuation of education (location/year), append to last education entry
                        if content['education']:
                            # Combine with comma separator
                            content['education'][-1] = f"{content['education'][-1]}, {clean_line}"
                    else:
                        # This is a certification entry
                        cert_clean = clean_line
                        cert_clean = re.sub(r'\.$', '', cert_clean).strip()
                        if cert_clean and len(cert_clean) > 3:
                            content['certifications'].append(cert_clean)
                # For education, handle "Degree | Institution | Year" format - only keep the last entry
                elif current_section == 'education':
                    # Clean up education entry
                    edu_clean = clean_line
                    # Remove any trailing periods or extra formatting
                    edu_clean = re.sub(r'\.$', '', edu_clean).strip()
                    # Ensure pipe-separated format if possible
                    if edu_clean and len(edu_clean) > 5 and edu_clean.lower() not in ['technical skills', 'industry experience', 'functional skills', 'education', 'certification', 'certifications', 'summary', 'projects experience']:
                        # Only keep the last education entry (replace previous ones)
                        content[current_section] = [edu_clean]
                elif clean_line and clean_line.lower() not in ['technical skills', 'industry experience', 'functional skills', 'education', 'certification', 'certifications', 'summary']:
                    content[current_section].append(clean_line)
            elif current_section == 'summary':
                # Summary paragraphs
                if line and not line.startswith('-') and not line.startswith('•'):
                    content['summary'].append(line)
            elif current_section == 'projects':
                # Enhanced project parsing - handle multiple formats:
                # Format 1: "Title / Company | Dates | Location" (en-dash format)
                # Format 2: "Title, Company, Location | Dates" (comma-separated format)
                line_clean = line.strip()
                
                # Check if line contains both title and dates (new single-line format)
                # Date patterns: MM/YYYY - MM/YYYY, YYYY-PRESENT, MM/YYYY – MM/YYYY
                date_pattern = r'([0-9]{1,2}/[0-9]{4}|[0-9]{4})\s*[-–—]\s*([0-9]{1,2}/[0-9]{4}|[0-9]{4}|present|current)'
                has_dates = re.search(date_pattern, line_clean, re.IGNORECASE)
                has_title_company = ' / ' in line_clean or (',' in line_clean and '|' in line_clean) or '|' in line_clean
                
                # If line has both title/company and dates, parse as single-line format
                if has_dates and has_title_company and not line.startswith('-') and not line.startswith('•'):
                    # Parse format: "Job Title / Company | Dates | Location" OR "Job Title, Company, Location | Dates"
                    parts = [p.strip() for p in line_clean.split('|')]
                    
                    if len(parts) >= 2:
                        # First part: Title / Company OR Title, Company, Location
                        title_company = parts[0].strip()
                        # Second part: Dates (remove "Dates: " prefix if present)
                        dates_text = re.sub(r'^dates?:\s*', '', parts[1], flags=re.IGNORECASE).strip()
                        # Third part (if exists and format 1): Location
                        location = parts[2].strip() if len(parts) > 2 else ''
                        
                        # If title_company contains commas, it might be format 2 (comma-separated)
                        # In format 2, location is already in title_company, so we don't need separate location
                        if ',' in title_company and not location:
                            # Format 2: "Title, Company, Location" - location is in the title part
                            # Keep as-is, location will be extracted from title if needed
                            pass
                        
                        # Save previous project if exists
                        if current_project and (current_project.get('title') or current_project.get('responsibilities')):
                            content['projects'].append(current_project)
                        
                        # Start new project
                        current_project = {
                            'title': title_company,
                            'dates': dates_text,
                            'location': location if location else '',
                            'responsibilities': [],
                            'technologies': []
                        }
                # Legacy format: separate lines for dates
                elif 'dates:' in line_lower or (has_dates and not has_title_company):
                    # Extract dates and location
                    dates_text = re.sub(r'dates?:', '', line, flags=re.IGNORECASE).strip()
                    # Try to extract location if present (after | or at the end)
                    location_match = re.search(r'\|\s*(.+)$', dates_text)
                    if location_match:
                        location = location_match.group(1).strip()
                        dates_text = re.sub(r'\s*\|\s*.+$', '', dates_text).strip()
                    else:
                        location = None
                    
                    if current_project:
                        current_project['dates'] = dates_text
                        if location:
                            current_project['location'] = location
                elif 'technologies:' in line_lower:
                    # Technologies line - improved extraction
                    tech_text = re.sub(r'technologies?:', '', line, flags=re.IGNORECASE).strip()
                    # Split by comma, semicolon, or "and"
                    techs = [t.strip() for t in re.split(r'[,;]|\s+and\s+', tech_text) if t.strip()]
                    if current_project:
                        current_project['technologies'] = techs
                elif line.startswith('-') or line.startswith('•') or line.startswith('▪'):
                    # Responsibilities/achievements with bullet points
                    clean_line = re.sub(r'^[-•*▪]\s*', '', line).strip()
                    if current_project and clean_line and len(clean_line) > 5:
                        # Skip if it looks like a section header
                        if not any(keyword in clean_line.lower() for keyword in ['technical skills', 'summary', 'education', 'certification', 'projects']):
                            current_project['responsibilities'].append(clean_line)
                elif line_clean.endswith(':') and not line_lower.startswith('technologies') and not line_lower.startswith('dates') and len(line_clean) < 100:
                    # This might be a category header (e.g., "ETL Development:", "Project Management:")
                    # Add it to responsibilities so it can be rendered as a category header
                    if current_project and line_clean and line_clean.count(':') == 1:
                        # Make sure it's not a main section header
                        if not any(keyword in line_lower for keyword in ['technical skills', 'industry experience', 'functional skills', 'education', 'summary', 'certification', 'projects experience', 'work experience']):
                            current_project['responsibilities'].append(line_clean)
                elif line_clean and not any(keyword in line_lower for keyword in ['left_column', 'right_column', 'projects experience', 'technical skills', 'industry experience', 'functional skills', 'education', 'summary', 'certification']):
                    # Check if it's a project title (not a date, not technologies, not a bullet)
                    if not re.match(r'^\d{1,2}/\d{4}', line_clean) and 'technologies' not in line_lower:
                        # Check if this line is a new project title (usually uppercase or has company indicators)
                        is_likely_title = (
                            line_clean.isupper() or  # All caps
                            ' / ' in line_clean or  # Has separator
                            '|' in line_clean or  # Has pipe separator
                            any(indicator in line_lower for indicator in [' – ', ' manager', ' consultant', ' lead', ' engineer', ' analyst', ' developer', ' director']) or  # Job title indicators
                            (current_project and len(current_project.get('responsibilities', [])) > 0)  # We already have responsibilities for current project
                        )
                        
                        if is_likely_title:
                            # This is a new project title
                            if current_project and (current_project.get('title') or current_project.get('responsibilities')):
                                # Save previous project
                                content['projects'].append(current_project)
                            # Start new project
                            current_project = {'title': line_clean, 'dates': '', 'location': '', 'responsibilities': [], 'technologies': []}
                        elif current_project:
                            # This is a responsibility without bullet point - treat it as unformatted text
                            # Split by sentence-ending punctuation if it's a long paragraph
                            if len(line_clean) > 150:
                                # Long paragraph - split into sentences
                                # Split by period, exclamation, or question mark followed by space and capital letter
                                sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', line_clean)
                                for sentence in sentences:
                                    sentence = sentence.strip()
                                    if sentence and len(sentence) > 10:
                                        current_project['responsibilities'].append(sentence)
                            else:
                                # Short text - add as single responsibility
                                if len(line_clean) > 10:
                                    current_project['responsibilities'].append(line_clean)
            
            i += 1
        
        # Add last project
        if current_project:
            content['projects'].append(current_project)
        
        return content
    
    async def _generate_standard_format(self, cv_text: str, output_path: str):
        """
        Generate standard format CV
        """
        doc = Document()
        
        # Set default font - paragraph size is 9pt
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(9)  # Paragraph font size set to 9pt
        
        # Parse and add content
        lines = cv_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect headings (lines that are short and might be section headers)
            is_heading = self._is_heading(line, lines)
            
            if is_heading:
                # Add heading
                p = doc.add_paragraph()
                p.add_run(line).bold = True
                p.add_run().font.size = Pt(14)
                p.paragraph_format.space_after = Pt(6)
            else:
                # Add regular paragraph
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(3)
        
        # Save document
        doc.save(output_path)
    
    def _is_heading(self, line: str, all_lines: List[str]) -> bool:
        """
        Determine if a line is a heading based on heuristics
        """
        # Common section headers
        section_keywords = [
            'contact', 'summary', 'objective', 'profile', 'experience',
            'education', 'skills', 'certifications', 'projects', 'publications',
            'awards', 'languages', 'references'
        ]
        
        line_lower = line.lower().strip()
        
        # Check if line matches section header patterns
        if any(keyword in line_lower for keyword in section_keywords):
            # Check if it's likely a header (short line, possibly all caps or title case)
            if len(line) < 50 and (line.isupper() or line.istitle()):
                return True
        
        # Check for markdown-style headers
        if line.startswith('#') or line.startswith('##') or line.startswith('---'):
            return True
        
        return False
    
    async def _generate_two_column_format(self, cv_text: str, output_path: str):
        """
        Generate CV in two-column format with proper bold formatting
        """
        doc = Document()
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Set default font - paragraph size is 9pt
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(9)  # Paragraph font size set to 9pt
        
        # Parse structured content
        content = self._parse_cv_content(cv_text)
        
        # Header section (full width)
        if content.get('header'):
            header_info = content['header']
            # Name
            if header_info.get('name'):
                name_p = doc.add_paragraph()
                name_run = name_p.add_run(header_info['name'])
                name_run.font.size = Pt(20)
                name_run.font.bold = True
                name_p.paragraph_format.space_after = Pt(3)
            
            # Title
            if header_info.get('title'):
                title_p = doc.add_paragraph()
                title_run = title_p.add_run(header_info['title'])
                title_run.font.size = Pt(12)
                title_run.font.bold = True
                title_p.paragraph_format.space_after = Pt(6)
            
            # Contact info
            if header_info.get('contact'):
                contact_p = doc.add_paragraph(header_info['contact'])
                contact_p.paragraph_format.space_after = Pt(12)
        
        # Create two-column layout using table
        table = doc.add_table(rows=1, cols=2)
        # Remove table borders for cleaner look
        table.style = None
        
        # Set column widths (40% left, 60% right)
        for row in table.rows:
            row.cells[0].width = Inches(2.5)
            row.cells[1].width = Inches(4.5)
            # Remove borders from cells
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
        
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]
        
        # LEFT COLUMN - Skills, Education, Certifications
        # Skills
        if content.get('skills'):
            self._add_section_header(left_cell, "SKILLS")
            for skill_category in content['skills']:
                # Skill category header (bold)
                cat_p = left_cell.add_paragraph()
                cat_run = cat_p.add_run(skill_category['category'] + ":")
                cat_run.font.bold = True
                cat_p.paragraph_format.space_after = Pt(2)
                
                # Skill items
                for skill in skill_category['items']:
                    self._add_bullet_point(left_cell, skill)
            self._add_spacing(left_cell)
        
        # Education
        if content.get('education'):
            self._add_section_header(left_cell, "EDUCATION")
            for edu in content['education']:
                edu_p = left_cell.add_paragraph()
                # Degree name (bold)
                if edu.get('degree'):
                    degree_run = edu_p.add_run(edu['degree'])
                    degree_run.font.bold = True
                # Institution and year
                if edu.get('institution') or edu.get('year'):
                    details = []
                    if edu.get('institution'):
                        details.append(edu['institution'])
                    if edu.get('year'):
                        details.append(edu['year'])
                    if details:
                        edu_p.add_run("\n" + " | ".join(details))
                edu_p.paragraph_format.space_after = Pt(6)
        
        # Certifications
        if content.get('certifications'):
            self._add_section_header(left_cell, "CERTIFICATIONS")
            for cert in content['certifications']:
                cert_p = left_cell.add_paragraph()
                cert_run = cert_p.add_run(cert)
                cert_run.font.bold = True
                cert_p.paragraph_format.space_after = Pt(3)
        
        # RIGHT COLUMN - Summary, Experience, Projects
        # Professional Summary
        if content.get('summary'):
            self._add_section_header(right_cell, "PROFESSIONAL SUMMARY")
            summary_p = right_cell.add_paragraph(content['summary'])
            summary_p.paragraph_format.space_after = Pt(12)
        
        # Experience
        if content.get('experience'):
            self._add_section_header(right_cell, "EXPERIENCE")
            for exp in content['experience']:
                # Job title (bold)
                title_p = right_cell.add_paragraph()
                title_run = title_p.add_run(exp.get('title', ''))
                title_run.font.bold = True
                title_run.font.size = Pt(12)
                
                # Company name (bold) and dates
                if exp.get('company') or exp.get('dates') or exp.get('location'):
                    title_p.add_run(" | ")
                    if exp.get('company'):
                        company_run = title_p.add_run(exp['company'])
                        company_run.font.bold = True
                    if exp.get('dates'):
                        if exp.get('company'):
                            title_p.add_run(" | ")
                        title_p.add_run(exp['dates'])
                    if exp.get('location'):
                        title_p.add_run(" | ")
                        title_p.add_run(exp['location'])
                title_p.paragraph_format.space_after = Pt(3)
                
                # Responsibilities
                if exp.get('responsibilities'):
                    for resp in exp['responsibilities']:
                        self._add_bullet_point(right_cell, resp)
                
                right_cell.add_paragraph()  # Spacing between experiences
        
        # Projects Experience (Full width after table)
        if content.get('projects'):
            doc.add_paragraph()  # Spacing
            projects_heading = doc.add_paragraph()
            projects_run = projects_heading.add_run("PROJECTS EXPERIENCE")
            projects_run.font.size = Pt(14)
            projects_run.font.bold = True
            projects_heading.paragraph_format.space_after = Pt(12)
            
            for project in content['projects']:
                # Project title (bold)
                if project.get('title'):
                    title_p = doc.add_paragraph()
                    title_run = title_p.add_run(project['title'])
                    title_run.font.bold = True
                    title_run.font.size = Pt(12)
                    title_p.paragraph_format.space_after = Pt(3)
                
                # Project description
                if project.get('description'):
                    desc_p = doc.add_paragraph(project['description'])
                    desc_p.paragraph_format.space_after = Pt(6)
                
                doc.add_paragraph()  # Spacing between projects
        
        doc.save(output_path)
    
    def _parse_cv_content(self, text: str) -> Dict:
        """Parse CV content from structured text"""
        content = {
            'header': {},
            'summary': '',
            'skills': [],
            'experience': [],
            'education': [],
            'certifications': [],
            'projects': []
        }
        
        lines = text.split('\n')
        current_section = None
        current_skill_category = None
        current_experience = None
        current_education = None
        current_project = None
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            line_lower = line.lower()
            
            # Section detection
            if line.startswith('[') and line.endswith(']'):
                section_name = line[1:-1].strip().upper()
                if section_name == 'HEADER':
                    current_section = 'header'
                elif 'SUMMARY' in section_name:
                    current_section = 'summary'
                elif section_name == 'SKILLS':
                    current_section = 'skills'
                elif section_name == 'EXPERIENCE':
                    current_section = 'experience'
                elif section_name == 'EDUCATION':
                    current_section = 'education'
                elif section_name == 'CERTIFICATIONS':
                    current_section = 'certifications'
                elif 'PROJECTS' in section_name:
                    current_section = 'projects'
                else:
                    current_section = None
                i += 1
                continue
            
            # Header parsing
            if current_section == 'header':
                # First non-empty line is usually name
                if not content['header'].get('name') and len(line) < 100 and not line.startswith('+') and '@' not in line and 'linkedin' not in line_lower and not line.startswith('['):
                    content['header']['name'] = line
                # Title line (usually after name, not contact info)
                elif content['header'].get('name') and not content['header'].get('title') and len(line) < 100 and not line.startswith('+') and '@' not in line and 'linkedin' not in line_lower:
                    content['header']['title'] = line
                # Contact info (contains phone, email, or links)
                elif '+' in line or '@' in line or 'linkedin' in line_lower:
                    # Clean up markdown links
                    contact_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)
                    if not content['header'].get('contact'):
                        content['header']['contact'] = contact_line
                    else:
                        content['header']['contact'] += " | " + contact_line
            
            # Summary parsing
            elif current_section == 'summary':
                if content['summary']:
                    content['summary'] += " " + line
                else:
                    content['summary'] = line
            
            # Skills parsing
            elif current_section == 'skills':
                # Check if it's a category header (ends with colon or is bold pattern)
                if ':' in line or line.startswith('- **') or line.startswith('**'):
                    # Extract category name
                    category_match = re.match(r'^[-*]?\s*\*\*([^*]+)\*\*:', line)
                    if category_match:
                        category_name = category_match.group(1).strip()
                    else:
                        category_name = line.replace(':', '').replace('**', '').replace('-', '').strip()
                    
                    if current_skill_category:
                        content['skills'].append(current_skill_category)
                    current_skill_category = {'category': category_name, 'items': []}
                elif current_skill_category and line.startswith('-'):
                    # Skill item
                    skill_item = re.sub(r'^-\s*', '', line).strip()
                    if skill_item:
                        current_skill_category['items'].append(skill_item)
            
            # Experience parsing
            elif current_section == 'experience':
                # Job title line (usually bold or starts with **)
                if line.startswith('**'):
                    # Format: **Job Title** | Company | Dates | Location
                    title_match = re.match(r'\*\*([^*]+)\*\*', line)
                    if title_match:
                        title = title_match.group(1).strip()
                        rest = line[title_match.end():].strip()
                        
                        if current_experience:
                            content['experience'].append(current_experience)
                        
                        current_experience = {
                            'title': title,
                            'company': '',
                            'dates': '',
                            'location': '',
                            'responsibilities': []
                        }
                        
                        # Parse company, dates, location from rest
                        if rest.startswith('|'):
                            parts = [p.strip() for p in rest[1:].split('|')]
                            if len(parts) >= 1:
                                current_experience['company'] = parts[0]
                            if len(parts) >= 2:
                                current_experience['dates'] = parts[1]
                            if len(parts) >= 3:
                                current_experience['location'] = parts[2]
                elif current_experience and '|' in line and not line.startswith('-') and not line.startswith('**'):
                    # Company, dates, location line (separate line)
                    parts = [p.strip() for p in line.split('|')]
                    # Check if this looks like dates (contains year pattern or "Present")
                    has_date_pattern = bool(re.search(r'\d{2}/\d{4}|Present|–|—', line))
                    
                    if has_date_pattern:
                        # This is a dates/location line
                        if len(parts) >= 1 and not current_experience.get('dates'):
                            current_experience['dates'] = parts[0]
                        if len(parts) >= 2 and not current_experience.get('location'):
                            current_experience['location'] = parts[1]
                    else:
                        # This might be company line
                        if len(parts) >= 1 and not current_experience.get('company'):
                            current_experience['company'] = parts[0]
                        if len(parts) >= 2 and not current_experience.get('dates'):
                            current_experience['dates'] = parts[1]
                        if len(parts) >= 3 and not current_experience.get('location'):
                            current_experience['location'] = parts[2]
                elif current_experience and line.startswith('-'):
                    # Responsibility
                    resp = re.sub(r'^-\s*', '', line).strip()
                    if resp:
                        current_experience['responsibilities'].append(resp)
            
            # Education parsing
            elif current_section == 'education':
                # Degree line (usually bold)
                if line.startswith('**') or (not line.startswith('-') and len(line) < 150):
                    if current_education:
                        content['education'].append(current_education)
                    
                    degree_match = re.match(r'\*\*([^*]+)\*\*', line)
                    if degree_match:
                        degree = degree_match.group(1).strip()
                        rest = line.replace(f'**{degree}**', '').strip()
                    else:
                        # Try to split degree and institution
                        if '|' in line:
                            parts = [p.strip() for p in line.split('|')]
                            degree = parts[0].replace('**', '').strip()
                            rest = ' | '.join(parts[1:]) if len(parts) > 1 else ''
                        else:
                            degree = line.replace('**', '').strip()
                            rest = ''
                    
                    current_education = {
                        'degree': degree,
                        'institution': '',
                        'year': ''
                    }
                    
                    # Parse institution and year from rest
                    if rest:
                        if '|' in rest:
                            parts = [p.strip() for p in rest.split('|')]
                            if len(parts) >= 1:
                                current_education['institution'] = parts[0]
                            if len(parts) >= 2:
                                current_education['year'] = parts[1]
                        else:
                            # Try to extract year
                            year_match = re.search(r'\b(19|20)\d{2}\b', rest)
                            if year_match:
                                current_education['year'] = year_match.group()
                                current_education['institution'] = rest.replace(year_match.group(), '').strip(' |,')
                            else:
                                current_education['institution'] = rest
            
            # Certifications parsing
            elif current_section == 'certifications':
                if line.startswith('-'):
                    cert = re.sub(r'^-\s*', '', line).strip()
                    if cert:
                        content['certifications'].append(cert)
                elif not line.startswith('-') and len(line) < 200:
                    # Might be a certification without bullet
                    content['certifications'].append(line)
            
            # Projects parsing
            elif current_section == 'projects':
                if line.startswith('**') or (not line.startswith('-') and len(line) < 100):
                    # Project title
                    if current_project:
                        content['projects'].append(current_project)
                    
                    title = line.replace('**', '').strip()
                    current_project = {'title': title, 'description': ''}
                elif current_project and line.startswith('-'):
                    # Project description item
                    desc = re.sub(r'^-\s*', '', line).strip()
                    if current_project['description']:
                        current_project['description'] += "\n" + desc
                    else:
                        current_project['description'] = desc
            
            i += 1
        
        # Add last items
        if current_skill_category:
            content['skills'].append(current_skill_category)
        if current_experience:
            content['experience'].append(current_experience)
        if current_education:
            content['education'].append(current_education)
        if current_project:
            content['projects'].append(current_project)
        
        return content

