import os
from docx import Document
from typing import Dict, Optional

# Try to import pdfplumber, fallback to pypdf if not available
try:
    import pdfplumber
    PDF_PARSER = 'pdfplumber'
except ImportError:
    try:
        from pypdf import PdfReader
        PDF_PARSER = 'pypdf'
    except ImportError:
        PDF_PARSER = None


class CVParser:
    """Parse CV from PDF or DOCX files"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    async def parse_cv(self, file_path: str) -> Optional[Dict[str, any]]:
        """
        Parse CV file and extract structured data
        Returns a dictionary with CV sections
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return await self._parse_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return await self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    async def _parse_pdf(self, file_path: str) -> Dict[str, any]:
        """Extract text from PDF"""
        text_content = []
        
        try:
            if PDF_PARSER == 'pdfplumber':
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
            elif PDF_PARSER == 'pypdf':
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            else:
                raise Exception("No PDF parser available. Please install pdfplumber or pypdf.")
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
        
        full_text = "\n\n".join(text_content)
        return self._structure_cv_data(full_text)
    
    async def _parse_docx(self, file_path: str) -> Dict[str, any]:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_content)
            return self._structure_cv_data(full_text)
        
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    def _structure_cv_data(self, text: str) -> Dict[str, any]:
        """
        Structure the extracted text into CV sections
        This is a basic implementation - can be enhanced with NLP
        """
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Basic structure - can be enhanced with more sophisticated parsing
        cv_data = {
            "raw_text": text,
            "full_text": "\n".join(lines),
            "sections": {
                "personal_info": [],
                "summary": [],
                "experience": [],
                "education": [],
                "skills": [],
                "certifications": [],
                "other": []
            }
        }
        
        # Simple keyword-based section detection
        current_section = "other"
        section_keywords = {
            "personal_info": ["name", "email", "phone", "address", "linkedin", "github"],
            "summary": ["summary", "objective", "profile", "about"],
            "experience": ["experience", "employment", "work history", "career"],
            "education": ["education", "academic", "qualifications", "degree"],
            "skills": ["skills", "competencies", "technical skills", "proficiencies"],
            "certifications": ["certifications", "certificates", "licenses"]
        }
        
        for line in lines:
            line_lower = line.lower()
            # Check if line is a section header
            for section, keywords in section_keywords.items():
                if any(keyword in line_lower for keyword in keywords) and len(line) < 50:
                    current_section = section
                    break
            
            if current_section in cv_data["sections"]:
                cv_data["sections"][current_section].append(line)
            else:
                cv_data["sections"]["other"].append(line)
        
        return cv_data

